import json
import logging
import re
from db import execute_write, execute_query
from llm_service import call_llm, load_prompt

def resolve_stakeholder_for_user(user_email, user_full_name, user_role):
    """Find an existing stakeholder matching this user's email, or create one."""
    existing = execute_query("SELECT id FROM stakeholders WHERE email = %s", (user_email,))
    if existing:
        return existing[0]['id']

    # Map the users.role value to the stakeholders.role ENUM
    role_map = {
        'leadership': 'leadership',
        'PwC Leadership': 'PwC Leadership',
        'engagement_manager': 'engagement_manager',
        'Delivery / Engagement Manager': 'Delivery / Engagement Manager',
        'manager': 'engagement_manager',
        'outgoing_sme': 'outgoing_sme',
        'Outgoing SME (Knowledge Giver)': 'Outgoing SME (Knowledge Giver)',
        'incoming_member': 'incoming_member',
        'Incoming Team Member (Knowledge Receiver)': 'Incoming Team Member (Knowledge Receiver)',
    }
    mapped_role = role_map.get(user_role, user_role)

    new_id = execute_write(
        "INSERT INTO stakeholders (name, email, role) VALUES (%s, %s, %s)",
        (user_full_name, user_email, mapped_role)
    )
    return new_id

def generate_plan_service(application_name, scope_description, plan_type, user_email=None, user_full_name=None, user_role=None, reverse_kt_focus=None, project_config=None, project_id=None):
    created_by = None
    if user_email and user_full_name and user_role:
        created_by = resolve_stakeholder_for_user(user_email, user_full_name, user_role)

    focus_text = f"\n    Reverse KT Focus Area: {reverse_kt_focus}" if reverse_kt_focus and plan_type == 'Reverse-KT' else ""
    project_config_text = f"\n    Project Configuration Details:\n{json.dumps(project_config, indent=2)}" if project_config else ""
    
    prompt = load_prompt(
        "plan_generation.txt",
        plan_type=plan_type,
        application_name=application_name,
        scope_description=scope_description,
        focus_text=focus_text,
        project_config_text=project_config_text
    )

    # Call LLM
    generated_content = call_llm(prompt)
    
    # Save to DB as draft
    query = """
        INSERT INTO kt_plans (application_name, scope_description, plan_type, generated_content, status, created_by, project_config, project_id)
        VALUES (%s, %s, %s, %s, 'draft', %s, %s, %s)
    """
    params = (application_name, scope_description, plan_type, generated_content, created_by, json.dumps(project_config) if project_config else None, project_id)
    plan_id = execute_write(query, params)
    
    # Extract topics
    extract_and_save_topics(plan_id, generated_content)
    
    return {
        "id": plan_id,
        "generated_content": generated_content,
        "status": "draft"
    }

def extract_and_save_topics(plan_id, generated_content):
    """Extract and save topics into plan_topics DB table using deterministic parser."""
    topics = parse_topics_directly_from_content(generated_content)
    execute_write("DELETE FROM plan_topics WHERE plan_id = %s", (plan_id,))
    count = 0
    for item in topics:
        query = "INSERT INTO plan_topics (plan_id, day_label, topic_name, estimated_duration_hours) VALUES (%s, %s, %s, %s)"
        execute_write(query, (plan_id, item.get('day_label', 'General'), item.get('topic_name'), item.get('estimated_duration_hours', 'N/A')))
        count += 1
    return count

def parse_topics_directly_from_content(generated_content):
    """Deterministic Python parser that reads markdown text directly to extract all sub-topic rows."""
    import re
    topics = []
    
    # Split content by Day headers or Post-KA H4 headers
    day_blocks = re.split(r'\n(?=Day\s+\d+:|####\s*Final\s*Assessment|####\s*Shadow\s*Resourcing|####\s*Lead\s*Resourcing)', generated_content, flags=re.IGNORECASE)
    
    ka_last_day = 1
    
    for block in day_blocks:
        block_strip = block.strip()
        day_match = re.match(r'^(Day\s+(\d+)[^\n]*)', block_strip, re.IGNORECASE)
        if not day_match:
            continue
            
        full_day_header = day_match.group(1).strip()
        day_num = int(day_match.group(2))
        
        # Skip if this is a post-KA phase header wrongly formatted as Day X
        if 'assessment' in block_strip.lower() or 'shadow' in block_strip.lower() or 'lead' in block_strip.lower():
            continue
            
        ka_last_day = max(ka_last_day, day_num)
        
        # Find all bullet lines: • Sub-topic Name (N minutes)
        bullet_matches = re.findall(r'[•\*]\s*([^\n\(]+?)\s*\((?:Part\s+\d+\s*-\s*)?(\d+)\s*minutes?\)', block_strip, re.IGNORECASE)
        
        if bullet_matches:
            for sub_name, mins_str in bullet_matches:
                sub_name_clean = sub_name.strip()
                if 'daily total' in sub_name_clean.lower() or 'timeline:' in sub_name_clean.lower():
                    continue
                mins = int(mins_str)
                hrs = round(mins / 60.0, 2)
                hrs_str = f"{hrs:g}"
                
                topics.append({
                    "day_label": full_day_header,
                    "topic_name": sub_name_clean,
                    "estimated_duration_hours": hrs_str
                })
        else:
            main_headings = re.findall(r'^\d+\.\s*([^\n]+)', block_strip, re.MULTILINE)
            for h_name in main_headings:
                h_clean = re.sub(r'\(continued\)', '', h_name, flags=re.IGNORECASE).strip()
                if h_clean and 'daily total' not in h_clean.lower() and 'timeline:' not in h_clean.lower():
                    topics.append({
                        "day_label": full_day_header,
                        "topic_name": h_clean,
                        "estimated_duration_hours": "2"
                    })

    # Post-KA Phases parsed directly from H4 headers or default
    asmt_match = re.search(r'####\s*Final\s*Assessment[\s\S]*?Timeline:\s*Day\s+(\d+)\s+to\s+Day\s+(\d+)\s*(?:\((\d+)\s*Days\))?', generated_content, re.IGNORECASE) or re.search(r'Timeline:\s*Day\s+(\d+)\s+to\s+Day\s+(\d+)\s*(?:\((\d+)\s*Days\))?', generated_content, re.IGNORECASE)
    sr_match = re.search(r'####\s*Shadow\s*Resourcing[\s\S]*?Timeline:\s*Day\s+(\d+)\s+to\s+Day\s+(\d+)', generated_content, re.IGNORECASE)
    lr_match = re.search(r'####\s*Lead\s*Resourcing[\s\S]*?Timeline:\s*Day\s+(\d+)\s+onwards', generated_content, re.IGNORECASE)

    if asmt_match:
        a_start, a_end = int(asmt_match.group(1)), int(asmt_match.group(2))
        a_days = int(asmt_match.group(3)) if (len(asmt_match.groups()) >= 3 and asmt_match.group(3)) else (a_end - a_start + 1)
        asmt_label = f"Day {a_start} to Day {a_end}"
        asmt_hrs = a_days * 24
        asmt_dur = f"{asmt_hrs} Hours ({a_days} Days)"
    else:
        a_days = 90
        a_start = ka_last_day + 1
        a_end = ka_last_day + a_days
        asmt_label = f"Day {a_start} to Day {a_end}"
        asmt_dur = f"{a_days * 24} Hours ({a_days} Days)"

    topics.append({
        "day_label": asmt_label,
        "topic_name": "Mandatory Final Assessment Evaluation Window",
        "estimated_duration_hours": asmt_dur
    })

    if sr_match:
        s_start, s_end = int(sr_match.group(1)), int(sr_match.group(2))
        sr_label = f"Day {s_start} to Day {s_end} (Shadow Phase)"
        s_days = max(1, s_end - s_start + 1)
        sr_dur = f"{s_days * 24} Hours ({s_days} Days / 2 Weeks)"
    else:
        s_days = 14
        s_start = a_end + 1
        s_end = s_start + 13
        sr_label = f"Day {s_start} to Day {s_end} (Shadow Phase)"
        sr_dur = f"{s_days * 24} Hours ({s_days} Days / 2 Weeks)"

    topics.append({
        "day_label": sr_label,
        "topic_name": "Practical Shadow Experience & Hands-on Ticket Resolution",
        "estimated_duration_hours": sr_dur
    })

    if lr_match:
        l_start = int(lr_match.group(1))
        lr_label = f"Day {l_start} onwards (Lead Phase)"
    else:
        l_start = s_end + 1
        lr_label = f"Day {l_start} onwards (Lead Phase)"

    topics.append({
        "day_label": lr_label,
        "topic_name": "Independent Project Leadership & Transition Completion",
        "estimated_duration_hours": "Ongoing (Lead Phase)"
    })

    return topics

def ensure_post_ka_phase_order(topics, generated_content, ka_last_day=1):
    """Separate KA topics and Post-KA topics, then reconstruct Post-KA topics strictly as: Assessment -> Shadow -> Lead."""
    import re
    
    ka_topics = []
    
    for t in topics:
        d_lbl = str(t.get('day_label', ''))
        t_nm = str(t.get('topic_name', ''))
        full_str = (d_lbl + ' ' + t_nm).lower()
        
        if 'assessment' in full_str or 'shadow' in full_str or 'lead' in full_str:
            continue
        ka_topics.append(t)
            
    # Parse exact timeline numbers from generated_content markdown strictly from H4 headers if present
    asmt_match = re.search(r'####\s*Final\s*Assessment[\s\S]*?Timeline:\s*Day\s+(\d+)\s+to\s+Day\s+(\d+)\s*(?:\((\d+)\s*Days\))?', generated_content, re.IGNORECASE) or re.search(r'Timeline:\s*Day\s+(\d+)\s+to\s+Day\s+(\d+)\s*(?:\((\d+)\s*Days\))?', generated_content, re.IGNORECASE)
    sr_match = re.search(r'####\s*Shadow\s*Resourcing[\s\S]*?Timeline:\s*Day\s+(\d+)\s+to\s+Day\s+(\d+)', generated_content, re.IGNORECASE)
    lr_match = re.search(r'####\s*Lead\s*Resourcing[\s\S]*?Timeline:\s*Day\s+(\d+)\s+onwards', generated_content, re.IGNORECASE)

    # 1. Assessment Topic
    if asmt_match:
        a_start, a_end = int(asmt_match.group(1)), int(asmt_match.group(2))
        a_days = int(asmt_match.group(3)) if (len(asmt_match.groups()) >= 3 and asmt_match.group(3)) else (a_end - a_start + 1)
        asmt_label = f"Day {a_start} to Day {a_end}"
        asmt_dur = f"{a_days * 24} Hours ({a_days} Days)"
    else:
        a_days = 90
        a_start = ka_last_day + 1
        a_end = ka_last_day + a_days
        asmt_label = f"Day {a_start} to Day {a_end}"
        asmt_dur = f"{a_days * 24} Hours ({a_days} Days)"
        
    asmt_topic = {
        "day_label": asmt_label,
        "topic_name": "Mandatory Final Assessment Evaluation Window",
        "estimated_duration_hours": asmt_dur
    }

    # 2. Shadow Topic
    if sr_match:
        s_start, s_end = int(sr_match.group(1)), int(sr_match.group(2))
        sr_label = f"Day {s_start} to Day {s_end} (Shadow Phase)"
        s_days = max(1, s_end - s_start + 1)
        sr_dur = f"{s_days * 24} Hours ({s_days} Days / 2 Weeks)"
    else:
        s_days = 14
        s_start = a_end + 1
        s_end = s_start + 13
        sr_label = f"Day {s_start} to Day {s_end} (Shadow Phase)"
        sr_dur = f"{s_days * 24} Hours ({s_days} Days / 2 Weeks)"
        
    sr_topic = {
        "day_label": sr_label,
        "topic_name": "Practical Shadow Experience & Hands-on Ticket Resolution",
        "estimated_duration_hours": sr_dur
    }

    # 3. Lead Topic
    if lr_match:
        l_start = int(lr_match.group(1))
        lr_label = f"Day {l_start} onwards (Lead Phase)"
    else:
        l_start = s_end + 1
        lr_label = f"Day {l_start} onwards (Lead Phase)"
        
    lr_topic = {
        "day_label": lr_label,
        "topic_name": "Independent Project Leadership & Transition Completion",
        "estimated_duration_hours": "Ongoing (Lead Phase)"
    }

    final_topics = ka_topics + [asmt_topic, sr_topic, lr_topic]
    return final_topics

def recalculate_plan_timeline_service(plan_id, new_assessment_days):
    """Dynamically update plan markdown text and topics in DB when manager saves new assessment window days."""
    try:
        plan_res = execute_query("SELECT generated_content, project_config FROM kt_plans WHERE id = %s", (plan_id,))
        if not plan_res or not plan_res[0].get('generated_content'):
            return
        
        content = plan_res[0]['generated_content']
        proj_config = plan_res[0].get('project_config')
        if proj_config and isinstance(proj_config, str):
            try:
                proj_config = json.loads(proj_config)
            except Exception:
                proj_config = {}
        elif not isinstance(proj_config, dict):
            proj_config = {}
            
        proj_config['final_deadline_extension_days'] = new_assessment_days

        # Dynamically determine the last KA day (ka_last_day) strictly from Knowledge Acquisition phase section before H4 Post-KA headers
        ka_content_part = re.split(r'####\s*(?:Final\s*Assessment|Shadow\s*Resourcing|Lead\s*Resourcing)', content, flags=re.IGNORECASE)[0]
        content_day_matches = re.findall(r'Day\s+(\d+)', ka_content_part, re.IGNORECASE)
        ka_last_day = max(int(m) for m in content_day_matches) if content_day_matches else 1

        asmt_start = ka_last_day + 1
        asmt_end = ka_last_day + new_assessment_days
        sr_start = asmt_end + 1
        sr_end = sr_start + 13 # 14 days
        lr_start = sr_end + 1

        # 1. Replace Assessment timeline in content under H4 header
        content = re.sub(
            r'(####\s*Final\s*Assessment[\s\S]*?)Timeline:\s*Day\s+\d+\s+to\s+Day\s+\d+(?:\s*\(\d+\s*Days\))?',
            rf'\g<1>Timeline: Day {asmt_start} to Day {asmt_end} ({new_assessment_days} Days)',
            content,
            count=1,
            flags=re.IGNORECASE
        )
        content = re.sub(
            r'((?:Mandatory\s*)?Final\s*Assessment\s*Evaluation\s*Window[^\n]*?:\s*)\d+(\s*Days)',
            rf'\g<1>{new_assessment_days}\g<2>',
            content,
            flags=re.IGNORECASE
        )

        # 2. Replace Shadow timeline in content under H4 header
        content = re.sub(
            r'(####\s*Shadow\s*Resourcing[\s\S]*?)Timeline:\s*Day\s+\d+\s+to\s+Day\s+\d+',
            rf'\g<1>Timeline: Day {sr_start} to Day {sr_end}',
            content,
            count=1,
            flags=re.IGNORECASE
        )

        # 3. Replace Lead timeline in content under H4 header
        content = re.sub(
            r'(####\s*Lead\s*Resourcing[\s\S]*?)Timeline:\s*Day\s+\d+\s+onwards',
            rf'\g<1>Timeline: Day {lr_start} onwards',
            content,
            count=1,
            flags=re.IGNORECASE
        )

        # Write updated content and project_config back to kt_plans table
        execute_write(
            "UPDATE kt_plans SET generated_content = %s, project_config = %s, final_deadline_extension_days = %s WHERE id = %s",
            (content, json.dumps(proj_config), new_assessment_days, plan_id)
        )

        # Re-sync topics table instantly using deterministic Python parser without slow LLM calls
        topics = parse_topics_directly_from_content(content)
        execute_write("DELETE FROM plan_topics WHERE plan_id = %s", (plan_id,))
        for item in topics:
            query = "INSERT INTO plan_topics (plan_id, day_label, topic_name, estimated_duration_hours) VALUES (%s, %s, %s, %s)"
            execute_write(query, (plan_id, item.get('day_label', 'General'), item.get('topic_name'), item.get('estimated_duration_hours', 'N/A')))

    except Exception as e:
        logging.error(f"Error in recalculate_plan_timeline_service: {e}")

def extract_and_save_topics(plan_id, generated_content):
    topics = []
    try:
        extraction_prompt = load_prompt("plan_topic_extraction.txt", generated_content=generated_content)
        extraction_response = call_llm(extraction_prompt)
        
        import re
        clean_json = extraction_response.strip()
        match = re.search(r'\[.*\]', clean_json, re.DOTALL)
        if match:
            clean_json = match.group(0)
            topics = json.loads(clean_json)
    except Exception as e:
        logging.warning(f"LLM topic extraction notice for plan {plan_id}: {e}")

    # Count extracted KA topics from LLM
    ka_topics_count = sum(1 for t in topics if 'assessment' not in (str(t.get('day_label','')) + str(t.get('topic_name',''))).lower() and 'shadow' not in (str(t.get('day_label','')) + str(t.get('topic_name',''))).lower() and 'lead' not in (str(t.get('day_label','')) + str(t.get('topic_name',''))).lower())

    # Count total Day X headings in generated_content markdown
    markdown_day_count = len(re.findall(r'^Day\s+\d+:', generated_content, re.MULTILINE | re.IGNORECASE))

    # If LLM response was truncated or yielded too few topics compared to markdown days, use deterministic Python parser
    if ka_topics_count < markdown_day_count:
        topics = parse_topics_directly_from_content(generated_content)
    else:
        # Dynamically determine the maximum KA day (ka_last_day) from extracted topics
        ka_last_day = 1
        filtered_topics = []
        for t in topics:
            d_label = str(t.get('day_label', ''))
            t_name = str(t.get('topic_name', ''))
            full_str = (d_label + ' ' + t_name).lower()
            
            # Skip any literal "Timeline:" bullet items accidentally extracted as topic rows
            if t_name.strip().lower().startswith('timeline:') or t_name.strip().lower().startswith('• timeline:'):
                continue

            # Skip any LLM-hallucinated daily 2-hour topics for shadow/ticket resolution in post-KA phase
            if ('shadow' in full_str or 'ticket resolution' in full_str or 'lead' in full_str) and re.search(r'Day\s+(?:1[7-9]|[2-9]\d)', d_label, re.IGNORECASE):
                continue

            filtered_topics.append(t)
            if 'assessment' not in full_str and 'shadow' not in full_str and 'lead' not in full_str:
                day_nums = re.findall(r'Day\s+(\d+)', d_label, re.IGNORECASE)
                for dn in day_nums:
                    ka_last_day = max(ka_last_day, int(dn))

        topics = filtered_topics

    # Reconstruct Post-KA topics strictly in order: Assessment -> Shadow -> Lead, matching markdown dates
    topics = ensure_post_ka_phase_order(topics, generated_content, ka_last_day if 'ka_last_day' in locals() else 1)

    # Clear existing topics if this is a resync
    execute_write("DELETE FROM plan_topics WHERE plan_id = %s", (plan_id,))
    
    count = 0
    for item in topics:
        query = "INSERT INTO plan_topics (plan_id, day_label, topic_name, estimated_duration_hours) VALUES (%s, %s, %s, %s)"
        execute_write(query, (plan_id, item.get('day_label', 'General'), item.get('topic_name'), item.get('estimated_duration_hours', 'N/A')))
        count += 1
        
    return count

def add_topic_service(plan_id, day_label, topic_name, estimated_duration_hours="N/A"):
    query = "INSERT INTO plan_topics (plan_id, day_label, topic_name, estimated_duration_hours) VALUES (%s, %s, %s, %s)"
    topic_id = execute_write(query, (plan_id, day_label or 'General', topic_name, estimated_duration_hours or 'N/A'))
    return topic_id

def update_topic_service(topic_id, day_label=None, topic_name=None, estimated_duration_hours=None):
    updates = []
    params = []
    if day_label is not None:
        updates.append("day_label = %s")
        params.append(day_label)
    if topic_name is not None:
        updates.append("topic_name = %s")
        params.append(topic_name)
    if estimated_duration_hours is not None:
        updates.append("estimated_duration_hours = %s")
        params.append(estimated_duration_hours)
    
    if not updates:
        return False
        
    params.append(topic_id)
    query = f"UPDATE plan_topics SET {', '.join(updates)} WHERE id = %s"
    execute_write(query, tuple(params))
    return True

def delete_topic_service(topic_id):
    query = "DELETE FROM plan_topics WHERE id = %s"
    execute_write(query, (topic_id,))
    return True

def extract_plan_info_from_doc_service(files_input):
    import os
    import re
    import json
    import logging

    if not isinstance(files_input, list):
        files_input = [files_input]

    combined_text = ""
    first_filename = ""

    for file_storage in files_input:
        filename = file_storage.filename or ""
        if not first_filename and filename:
            first_filename = filename
        ext = os.path.splitext(filename)[1].lower()
        doc_text = ""

        if ext == '.pdf':
            try:
                import pypdf
                pdf_reader = pypdf.PdfReader(file_storage)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        doc_text += page_text + "\n"
            except Exception as e:
                logging.error(f"Error reading PDF {filename}: {e}")
                raise Exception(f"Failed to parse PDF file ({filename}): {str(e)}")
        elif ext in ['.docx', '.doc', '.docs']:
            try:
                import docx
                doc = docx.Document(file_storage)
                for para in doc.paragraphs:
                    if para.text and para.text.strip():
                        doc_text += para.text.strip() + "\n"
                for table in doc.tables:
                    for row in table.rows:
                        row_cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                        if row_cells:
                            doc_text += " | ".join(row_cells) + "\n"
            except Exception as e:
                logging.error(f"Error reading DOC/DOCX {filename}: {e}")
                raise Exception(f"Failed to parse Word document ({filename}): {str(e)}")
        elif ext == '.txt':
            doc_text = file_storage.read().decode('utf-8', errors='ignore')
        else:
            raise Exception(f"Unsupported file type for {filename}. Only PDF (.pdf) and Word documents (.doc, .docx) are allowed.")

        doc_text = doc_text.strip()
        if doc_text:
            combined_text += f"\n--- Document: {filename} ---\n{doc_text}\n"

    combined_text = combined_text.strip()
    if not combined_text:
        raise Exception("The uploaded document(s) are empty or text could not be extracted.")

    prompt = load_prompt(
        "plan_document_analysis.txt",
        file_count=len(files_input),
        combined_text=combined_text[:35000]
    )

    llm_res = call_llm(prompt)
    extracted_app_name = ""
    extracted_scope = ""

    if llm_res and isinstance(llm_res, str):
        match = re.search(r'\{.*\}', llm_res, re.DOTALL)
        if match:
            try:
                extracted = json.loads(match.group(0))
                extracted_app_name = str(extracted.get("application_name", "")).strip()
                extracted_scope = str(extracted.get("scope_description", "")).strip()
            except Exception as e:
                logging.error(f"Failed to parse LLM json regex match: {e}")

    # Helper function to clean application name
    def clean_name(raw_name):
        if not raw_name:
            return ""
        cleaned = re.sub(r'^(file|doc|document)\s*\d*[:\s-]*', '', raw_name, flags=re.IGNORECASE)
        cleaned = re.sub(r'\.(pdf|docx?|txt)$', '', cleaned, flags=re.IGNORECASE)
        cleaned = cleaned.replace('_', ' ').replace('-', ' ').strip()
        return cleaned.title()

    if not extracted_app_name or re.match(r'^(file|doc|document)\s*\d*', extracted_app_name, re.IGNORECASE):
        proj_match = re.search(r'(?:Project|System|Application)\s*Name[:\s-]+([^\n\r]+)', combined_text, re.IGNORECASE)
        if proj_match:
            extracted_app_name = clean_name(proj_match.group(1).strip())
        elif first_filename:
            extracted_app_name = clean_name(first_filename)
        else:
            extracted_app_name = "Hospital Management System"
    else:
        extracted_app_name = clean_name(extracted_app_name)

    if not extracted_scope or extracted_scope.startswith("--- Document:"):
        clean_lines = re.sub(r'---\s*Document:[^\n]+\n?', '', combined_text)
        clean_lines = re.sub(r'File\s*\d+:[^\n]+\n?', '', clean_lines)
        lines = [l.strip() for l in clean_lines.splitlines() if l.strip()]
        
        topics = []
        skip_words = {'contents', 'project name', 'project overview', 'project scope', 'functional modules', 'table of contents'}
        for line in lines:
            c_line = re.sub(r'^\d+[\.\)]\s*', '', line)
            c_line = re.sub(r'^[\-\*•]\s*', '', c_line).strip()
            if c_line and len(c_line) < 80 and c_line.lower() not in skip_words:
                if c_line not in topics and not c_line.lower().startswith('---'):
                    topics.append(c_line)
        if topics:
            extracted_scope = ", ".join(topics[:25])
        else:
            extracted_scope = re.sub(r'\s+', ' ', clean_lines[:500]).strip()

    return {
        "application_name": extracted_app_name or "Hospital Management System",
        "scope_description": extracted_scope
    }

def recalculate_plan_timeline_service(plan_id, new_assessment_days):
    """Dynamically update plan markdown text and topics in DB when manager saves new assessment window days."""
    try:
        plan_res = execute_query("SELECT generated_content, project_config FROM kt_plans WHERE id = %s", (plan_id,))
        if not plan_res or not plan_res[0].get('generated_content'):
            return
        
        content = plan_res[0]['generated_content']
        proj_config = plan_res[0].get('project_config')
        if proj_config and isinstance(proj_config, str):
            try:
                proj_config = json.loads(proj_config)
            except Exception:
                proj_config = {}
        elif not isinstance(proj_config, dict):
            proj_config = {}
            
        proj_config['final_deadline_extension_days'] = new_assessment_days

        # Dynamically determine ka_last_day strictly from the KA phase section (before Final Assessment header)
        ka_content_part = re.split(r'####?\s*Final\s*Assessment', content, flags=re.IGNORECASE)[0]
        content_day_matches = re.findall(r'Day\s+(\d+)', ka_content_part, re.IGNORECASE)
        ka_last_day = max(int(m) for m in content_day_matches) if content_day_matches else 10

        asmt_start = ka_last_day + 1
        asmt_end = ka_last_day + new_assessment_days
        sr_start = asmt_end + 1
        sr_end = sr_start + 13 # 14 days
        lr_start = sr_end + 1

        # Reconstruct the 3 Post-KA Phase blocks cleanly and unambiguously
        # 1. Update Final Assessment Window block
        def update_asmt_block(m):
            header = m.group(1)
            block = m.group(2)
            block = re.sub(
                r'([•\*]?\s*Timeline:\s*)Day\s+\d+\s+to\s+Day\s+\d+(?:\s*\([^)]*\))?',
                rf'\g<1>Day {asmt_start} to Day {asmt_end} ({new_assessment_days} Days)',
                block,
                flags=re.IGNORECASE
            )
            block = re.sub(
                r'((?:Mandatory\s*)?Final\s*Assessment\s*Evaluation\s*Window[^\n]*?:\s*)\d+(\s*Days)',
                rf'\g<1>{new_assessment_days}\g<2>',
                block,
                flags=re.IGNORECASE
            )
            return header + block

        content = re.sub(
            r'(####?\s*Final\s*Assessment[^\n]*\n)([\s\S]*?)(?=\n####?\s*Shadow|\n####?\s*Lead|\n##|\Z)',
            update_asmt_block,
            content,
            flags=re.IGNORECASE
        )

        # 2. Update Shadow Resourcing Phase block
        def update_sr_block(m):
            header = m.group(1)
            block = m.group(2)
            block = re.sub(
                r'([•\*]?\s*Timeline:\s*)Day\s+\d+\s+to\s+Day\s+\d+(?:\s*\([^)]*\))?',
                rf'\g<1>Day {sr_start} to Day {sr_end} (Week 1-2: 14 Days / 2 Weeks)',
                block,
                flags=re.IGNORECASE
            )
            return header + block

        content = re.sub(
            r'(####?\s*Shadow\s*Resourcing[^\n]*\n)([\s\S]*?)(?=\n####?\s*Lead|\n##|\Z)',
            update_sr_block,
            content,
            flags=re.IGNORECASE
        )

        # 3. Update Lead Resourcing Phase block
        def update_lr_block(m):
            header = m.group(1)
            block = m.group(2)
            block = re.sub(
                r'([•\*]?\s*Timeline:\s*)Day\s+\d+\s+onwards',
                rf'\g<1>Day {lr_start} onwards',
                block,
                flags=re.IGNORECASE
            )
            return header + block

        content = re.sub(
            r'(####?\s*Lead\s*Resourcing[^\n]*\n)([\s\S]*?)(?=\n##|\Z)',
            update_lr_block,
            content,
            flags=re.IGNORECASE
        )

        # Write updated content and project_config back to kt_plans table
        execute_write(
            "UPDATE kt_plans SET generated_content = %s, project_config = %s, final_deadline_extension_days = %s WHERE id = %s",
            (content, json.dumps(proj_config), new_assessment_days, plan_id)
        )

        # Re-sync topics table instantly using deterministic Python parser without slow LLM calls
        topics = parse_topics_directly_from_content(content)
        execute_write("DELETE FROM plan_topics WHERE plan_id = %s", (plan_id,))
        for item in topics:
            query = "INSERT INTO plan_topics (plan_id, day_label, topic_name, estimated_duration_hours) VALUES (%s, %s, %s, %s)"
            execute_write(query, (plan_id, item.get('day_label', 'General'), item.get('topic_name'), item.get('estimated_duration_hours', 'N/A')))

    except Exception as e:
        logging.error(f"Error in recalculate_plan_timeline_service: {e}")
