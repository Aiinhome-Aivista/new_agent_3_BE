from db import execute_query, execute_write
from llm_service import call_llm, load_prompt
import json
import os
from datetime import datetime
try:
    from docx import Document
except ImportError:
    Document = None

REPORTS_DIR = os.path.join(os.path.dirname(__file__), '..', 'reports_output')
os.makedirs(REPORTS_DIR, exist_ok=True)

def detect_risks_service(plan_id):
    # Gather data for LLM
    plan_query = "SELECT application_name, scope_description FROM kt_plans WHERE id = %s"
    plan_data = execute_query(plan_query, (plan_id,))
    
    comp_query = "SELECT topic, completion_percent FROM completion_tracking WHERE plan_id = %s"
    comp_data = execute_query(comp_query, (plan_id,))
    
    att_query = """
        SELECT m.title as topic_title, s.name as stakeholder_name, s.role, a.attended 
        FROM attendance a
        JOIN meetings m ON a.meeting_id = m.id
        JOIN stakeholders s ON a.stakeholder_id = s.id
        WHERE m.plan_id = %s
    """
    detailed_att_data = execute_query(att_query, (plan_id,))
    
    assess_query = """
        SELECT question, answer, s.name as stakeholder_name
        FROM assessments a
        JOIN stakeholders s ON a.stakeholder_id = s.id
        WHERE a.plan_id = %s
    """
    assess_data = execute_query(assess_query, (plan_id,))
    
    from rag_service import query_knowledge
    # Query with broader terms and more results to capture document context
    rag_chunks = query_knowledge("risks issues problems challenges gaps delays", plan_id, n_results=10)
    from guardrails import retrieval_rail
    # Relax threshold for L2 distance to ensure documents are not wrongly discarded
    retrieval_passed, _ = retrieval_rail(rag_chunks, threshold=1.5, endpoint="/api/risks/detect")
    if not retrieval_passed:
        rag_chunks = []
    rag_context = "\n".join([chunk["text"] for chunk in rag_chunks]) if rag_chunks else "None"
    
    prompt = load_prompt(
        "risk_detection.txt",
        plan_info=(plan_data[0] if plan_data else 'N/A'),
        comp_data=comp_data,
        detailed_att_data=detailed_att_data,
        assess_data=assess_data,
        rag_context=rag_context
    )
    
    llm_response = call_llm(prompt)
    
    # Try to parse JSON from LLM
    try:
        # Strip markdown code blocks if any
        clean_json = llm_response.replace('```json', '').replace('```', '').strip()
        risks = json.loads(clean_json)
    except json.JSONDecodeError:
        import logging
        logging.warning(f"Risk detection LLM parse failure for plan {plan_id}: {llm_response[:200]}")
        risks = []
        
    # === NEW RULE-BASED RISKS FOR 100% COMPLETED PLANS ===
    is_100_percent = len(comp_data) > 0 and all(c['completion_percent'] == 100 for c in comp_data)
    if is_100_percent:
        try:
            # 1. Fetch config and check options
            proj_config_query = "SELECT p.config FROM kt_projects p JOIN kt_plans pl ON pl.project_id = p.id WHERE pl.id = %s"
            proj_config_data = execute_query(proj_config_query, (plan_id,))
            if proj_config_data and proj_config_data[0]['config']:
                import json
                config_json = json.loads(proj_config_data[0]['config'])
                app_name = plan_data[0]['application_name'] if plan_data else ''
                track = next((t for t in config_json.get('tracks', []) if str(t.get('name', '')).strip() == str(app_name).strip()), None)
                
                if track and track.get('options'):
                    options = track['options']
                    
                    # 2. SUD Document Risk Check
                    if options.get('sud_doc_upload') or options.get('upload_sud'):
                        sud_req_query = "SELECT stakeholder_id FROM resource_mapping WHERE plan_id = %s AND is_sudo = 1"
                        sud_req_sh = execute_query(sud_req_query, (plan_id,))
                        for sh in sud_req_sh:
                            sh_id = sh['stakeholder_id']
                            sud_doc_check = execute_query("SELECT id FROM sud_documents WHERE plan_id = %s AND stakeholder_id = %s", (plan_id, sh_id))
                            if not sud_doc_check:
                                sh_name_data = execute_query("SELECT name FROM stakeholders WHERE id = %s", (sh_id,))
                                sh_name = sh_name_data[0]['name'] if sh_name_data else "Unknown"
                                risks.append({
                                    "description": f"Participant {sh_name} has not uploaded the required SUD document for a completed plan.",
                                    "severity": "high"
                                })
                                
                    # 3. Final Assessment Risk Check
                    if options.get('assessment') or options.get('assessment_80'):
                        assess_req_query = "SELECT stakeholder_id FROM resource_mapping WHERE plan_id = %s AND is_final_assessment = 1"
                        assess_req_sh = execute_query(assess_req_query, (plan_id,))
                        
                        plan_meta = execute_query("SELECT unlocked_on, final_deadline_extension_days FROM kt_plans WHERE id = %s", (plan_id,))
                        if plan_meta:
                            unlocked_on = plan_meta[0]['unlocked_on']
                            ext_days = plan_meta[0]['final_deadline_extension_days'] or 90
                            
                            from datetime import datetime, timedelta
                            deadline_missed = False
                            if unlocked_on:
                                deadline_date = unlocked_on + timedelta(days=ext_days)
                                if datetime.now() > deadline_date:
                                    deadline_missed = True
                            
                            for sh in assess_req_sh:
                                sh_id = sh['stakeholder_id']
                                assess_check = execute_query("SELECT id FROM assessment_results WHERE plan_id = %s AND stakeholder_id = %s AND assessment_type = 'final'", (plan_id, sh_id))
                                if not assess_check:
                                    sh_name_data = execute_query("SELECT name FROM stakeholders WHERE id = %s", (sh_id,))
                                    sh_name = sh_name_data[0]['name'] if sh_name_data else "Unknown"
                                    
                                    desc = f"Participant {sh_name} has not given the final assessment. The deadline of {ext_days} days is at risk."
                                    if deadline_missed:
                                        desc = f"Participant {sh_name} has missed the final assessment deadline."
                                    
                                    risks.append({
                                        "description": desc,
                                        "severity": "high"
                                    })
        except Exception as ex:
            import logging
            logging.error(f"Rule-based risk evaluation failed: {str(ex)}")
    # =========================================================

    saved_risks = []
    
    existing_query = "SELECT id, description, severity, status, detected_by FROM risks WHERE plan_id = %s AND status = 'open'"
    existing_risks = execute_query(existing_query, (plan_id,))
    
    def clean_words(s):
        return set(''.join(c if c.isalnum() or c.isspace() else ' ' for c in s.lower()).split())
        
    from guardrails import execution_rail
    for risk in risks:
        desc = risk.get('description', 'Unknown risk')
        severity = risk.get('severity', 'medium').lower()
        
        exec_passed, _ = execution_rail("risk_severity", {"severity": severity}, "/api/risks/detect")
        if not exec_passed:
            severity = 'medium'
            
        desc_words = clean_words(desc)
        duplicate_found = False
        
        for er in existing_risks:
            er_words = clean_words(er['description'])
            if not desc_words or not er_words:
                continue
            intersection = desc_words.intersection(er_words)
            if len(intersection) / len(desc_words) >= 0.6 or len(intersection) / len(er_words) >= 0.6:
                duplicate_found = True
                saved_risks.append({
                    "id": er['id'],
                    "description": er['description'],
                    "severity": er['severity'],
                    "status": er['status'],
                    "detected_by": er['detected_by']
                })
                break
                
        if duplicate_found:
            continue
            
        query = """
            INSERT INTO risks (plan_id, description, severity, detected_by)
            VALUES (%s, %s, %s, 'ai')
        """
        risk_id = execute_write(query, (plan_id, desc, severity))
            
        saved_risks.append({
            "id": risk_id,
            "description": desc,
            "severity": severity,
            "status": "open",
            "detected_by": "ai"
        })
        
    return saved_risks

def escalate_risk_service(risk_id, assigned_to=[], initial_note=None, manager_id=None):
    risk_query = "SELECT description, severity, plan_id FROM risks WHERE id = %s"
    risk_data = execute_query(risk_query, (risk_id,))
    if not risk_data:
        raise Exception("Risk not found")
        
    desc = risk_data[0]['description']
    severity = risk_data[0]['severity']
    plan_id = risk_data[0]['plan_id']
    
    from connectors import JiraConnector
    jira = JiraConnector()
    jira_res = jira.push_risk_to_jira(desc, severity, plan_id)
    jira_ref = jira_res.get("issue_id", "Unknown") if isinstance(jira_res, dict) else str(jira_res)
    
    query = "UPDATE risks SET jira_ticket_ref = %s WHERE id = %s"
    execute_write(query, (jira_ref, risk_id))
    
    del_query = "DELETE FROM risk_assignments WHERE risk_id = %s"
    execute_write(del_query, (risk_id,))
    
    for stakeholder_id in assigned_to:
        assign_query = "INSERT INTO risk_assignments (risk_id, stakeholder_id) VALUES (%s, %s)"
        execute_write(assign_query, (risk_id, stakeholder_id))
        
    if initial_note and manager_id:
        comment_query = "INSERT INTO risk_comments (risk_id, stakeholder_id, comment_text) VALUES (%s, %s, %s)"
        execute_write(comment_query, (risk_id, manager_id, initial_note))
    
    return {"escalated": True, "jira_ticket_ref": jira_ref, "assigned_to": assigned_to}

def get_assigned_risks_service(stakeholder_id):
    query = """
        SELECT DISTINCT r.*, p.application_name as plan_name 
        FROM risks r
        JOIN risk_assignments ra ON r.id = ra.risk_id
        JOIN kt_plans p ON r.plan_id = p.id
        WHERE ra.stakeholder_id = %s
        ORDER BY r.created_at DESC
    """
    risks = execute_query(query, (stakeholder_id,))
    
    for r in risks:
        comments_query = """
            SELECT rc.id, rc.comment_text, rc.created_at, s.name as stakeholder_name, s.role 
            FROM risk_comments rc
            JOIN stakeholders s ON rc.stakeholder_id = s.id
            WHERE rc.risk_id = %s
            ORDER BY rc.created_at ASC
        """
        r['comments'] = execute_query(comments_query, (r['id'],))
        
    return risks

def add_risk_comment_service(risk_id, stakeholder_id, comment_text):
    risk_query = "SELECT status FROM risks WHERE id = %s"
    risk = execute_query(risk_query, (risk_id,))
    if not risk:
        raise Exception("Risk not found")
    if risk[0]['status'] == 'solved':
        raise Exception("Cannot add comment to a solved risk")
        
    query = "INSERT INTO risk_comments (risk_id, stakeholder_id, comment_text) VALUES (%s, %s, %s)"
    comment_id = execute_write(query, (risk_id, stakeholder_id, comment_text))
    return {"id": comment_id, "comment_text": comment_text}

def update_risk_status_service(risk_id, status):
    risk_query = "SELECT status FROM risks WHERE id = %s"
    risk = execute_query(risk_query, (risk_id,))
    if not risk:
        raise Exception("Risk not found")
    if risk[0]['status'] == 'solved':
        raise Exception("Cannot change status of a solved risk")
        
    query = "UPDATE risks SET status = %s WHERE id = %s"
    execute_write(query, (status, risk_id))
    return True

def sort_risks_ui_order(risks):
    def get_risk_day(r):
        val = r.get('created_at') or r.get('date')
        if isinstance(val, datetime):
            return val.date()
        elif hasattr(val, 'date') and callable(val.date):
            return val.date()
        elif isinstance(val, str):
            try:
                return datetime.fromisoformat(str(val).replace('Z', '')).date()
            except Exception:
                return datetime.min.date()
        return datetime.min.date()

    def get_severity_rank(r):
        sev = str(r.get('severity', '')).lower()
        mapping = {'critical': 1, 'high': 2, 'medium': 3, 'low': 4}
        return mapping.get(sev, 5)

    return sorted(risks, key=lambda r: (
        1 if str(r.get('status', '')).lower() in ['solved', 'resolved', 'approved'] else 0,
        -get_risk_day(r).toordinal(),
        get_severity_rank(r)
    ))

def generate_risks_word_doc(plan_id):
    if not Document:
        raise Exception("python-docx is not installed")
        
    plan_query = "SELECT application_name, scope_description FROM kt_plans WHERE id = %s"
    plan_res = execute_query(plan_query, (plan_id,))
    plan_name = plan_res[0]['application_name'] if plan_res else f"Plan {plan_id}"
    scope_desc = plan_res[0]['scope_description'] if plan_res and 'scope_description' in plan_res[0] else ""
    
    query = "SELECT * FROM risks WHERE plan_id = %s"
    risks = execute_query(query, (plan_id,))
    risks = sort_risks_ui_order(risks)
    
    doc = Document()
    doc.add_heading(f"Risk Log & Details: {plan_name}", level=0)
    
    p_meta = doc.add_paragraph()
    p_meta.add_run("Plan ID: ").bold = True
    p_meta.add_run(f"{plan_id}    |    ")
    p_meta.add_run("Total Risks Logged: ").bold = True
    p_meta.add_run(f"{len(risks)}    |    ")
    p_meta.add_run("Generated At: ").bold = True
    p_meta.add_run(datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    
    if scope_desc:
        p_scope = doc.add_paragraph()
        p_scope.add_run("Scope / Description: ").bold = True
        p_scope.add_run(str(scope_desc))
        
    if not risks:
        doc.add_paragraph("No risks have been logged for this plan yet.")
    else:
        for idx, r in enumerate(risks, 1):
            title = r.get('description', 'Untitled Risk')
            short_title = title[:80] + "..." if len(title) > 80 else title
            doc.add_heading(f"{idx}. {short_title} (ID #{r['id']})", level=1)
            
            assignees_query = """
                SELECT s.name 
                FROM risk_assignments ra
                JOIN stakeholders s ON ra.stakeholder_id = s.id
                WHERE ra.risk_id = %s
            """
            assignees = execute_query(assignees_query, (r['id'],))
            assigned_names = [a['name'] for a in assignees] if assignees else []
            
            comments_query = """
                SELECT rc.comment_text, rc.created_at, s.name as stakeholder_name, s.role 
                FROM risk_comments rc
                JOIN stakeholders s ON rc.stakeholder_id = s.id
                WHERE rc.risk_id = %s
                ORDER BY rc.created_at ASC
            """
            comments = execute_query(comments_query, (r['id'],))
            
            def add_field_line(doc_obj, label, value):
                p = doc_obj.add_paragraph()
                p.paragraph_format.space_after = 2
                p.add_run(f"{label}: ").bold = True
                p.add_run(str(value))
                
            add_field_line(doc, "Risk ID", f"#{r['id']}")
            add_field_line(doc, "Full Description", r.get('description', 'N/A'))
            add_field_line(doc, "Severity", str(r.get('severity', '')).upper())
            add_field_line(doc, "Status", str(r.get('status', '')).upper())
            add_field_line(doc, "Detected By", str(r.get('detected_by', 'AI')).upper())
            add_field_line(doc, "Created At", str(r.get('created_at', 'N/A')))
            add_field_line(doc, "Assigned Stakeholders", ", ".join(assigned_names) if assigned_names else "Unassigned")
            
            doc.add_heading("Comments & Action History", level=2)
            if comments:
                for c in comments:
                    p_c = doc.add_paragraph(style='List Bullet')
                    author = c.get('stakeholder_name', 'Unknown')
                    role = c.get('role', '')
                    timestamp = str(c.get('created_at', ''))
                    p_c.add_run(f"[{timestamp}] {author} ({role}): ").bold = True
                    p_c.add_run(str(c.get('comment_text', '')))
            else:
                doc.add_paragraph("No comments or status updates recorded.")
                
    safe_plan_name = "".join([c if c.isalnum() else "_" for c in plan_name])
    filename = f"Risks_{safe_plan_name}_{plan_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join(REPORTS_DIR, filename)
    doc.save(filepath)
    return {"filename": filename, "filepath": filepath}

