from docx import Document

doc = Document()
doc.add_heading('Sample Project Requirement Document', 0)

doc.add_heading('Project Overview', level=1)
doc.add_paragraph('This document outlines the requirements for the newly proposed "SmartInventory" system. '
                  'The system will provide real-time tracking of warehouse inventory, supplier management, '
                  'and automated restocking alerts.')

doc.add_heading('Key Objectives', level=1)
doc.add_paragraph('1. Reduce manual inventory checks by 80%.\n'
                  '2. Automate supplier purchase orders when stock is below 10%.\n'
                  '3. Provide a dashboard for warehouse managers.')

doc.add_heading('Technical Stack', level=1)
doc.add_paragraph('Frontend: React, TailwindCSS\n'
                  'Backend: Python, Flask\n'
                  'Database: MySQL')

doc.save('sample_project.docx')
print("Created sample_project.docx successfully.")
