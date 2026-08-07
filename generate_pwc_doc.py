from docx import Document

doc = Document()

# Add Title
doc.add_heading('PwC Financial Transformation - Knowledge Transition (KT) Document', 0)

doc.add_heading('1. Project Overview', level=1)
doc.add_paragraph(
    'This project focuses on the digital transformation of financial workflows for a Fortune 500 client. '
    'The core objective is to migrate legacy on-premise ERP systems to a cloud-based SAP S/4HANA environment '
    'while implementing automated invoice processing using AI/OCR models. '
    'The incoming team members must understand the end-to-end architecture, API integrations, and the '
    'governance model established by PwC.'
)

doc.add_heading('2. Scope of Transition', level=1)
doc.add_paragraph(
    'The transition is divided into several modules, requiring deep technical understanding and operational readiness:'
)
doc.add_paragraph('Module 1: Legacy ERP Data Migration Strategies', style='List Bullet')
doc.add_paragraph('Module 2: Cloud Infrastructure & Security Governance', style='List Bullet')
doc.add_paragraph('Module 3: AI/OCR Invoice Automation Pipeline', style='List Bullet')
doc.add_paragraph('Module 4: API Integrations and Webhooks', style='List Bullet')

doc.add_heading('3. Key Technical Stack', level=1)
doc.add_paragraph(
    '- Backend: Python, FastAPI (for custom microservices)\n'
    '- Cloud: Microsoft Azure (AKS, Blob Storage)\n'
    '- Database: Azure PostgreSQL, SAP HANA\n'
    '- DevOps: Azure DevOps, Terraform\n'
    '- Monitoring: Datadog, Splunk'
)

doc.add_heading('4. Transition Goals', level=1)
doc.add_paragraph(
    'The incoming team is expected to fully take over the development and Level 3 support of the platform. '
    'A successful transition requires not just theoretical knowledge but practical shadow experience '
    'and hands-on ticket resolution to ensure SLA adherence.'
)

doc.save('PwC_Real_Project_Requirement.docx')
print("Created PwC_Real_Project_Requirement.docx successfully.")
